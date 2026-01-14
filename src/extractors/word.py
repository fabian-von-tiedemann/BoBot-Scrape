"""
Word document text extraction using python-docx.
"""
import logging
from typing import Optional

from docx import Document

logger = logging.getLogger(__name__)


def extract_word(filepath: str) -> Optional[str]:
    """
    Extract text from a Word document (.docx).

    Args:
        filepath: Path to the .docx file

    Returns:
        Extracted text content, or None if extraction fails
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        logger.warning(f"Failed to open Word document '{filepath}': {e}")
        return None

    try:
        parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    parts.append(" | ".join(row_text))

        if not parts:
            logger.warning(f"No text extracted from Word document '{filepath}'")
            return None

        return "\n\n".join(parts)

    except Exception as e:
        logger.warning(f"Error processing Word document '{filepath}': {e}")
        return None
